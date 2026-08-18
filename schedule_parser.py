#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Парсер расписания ветеринарной академии
Загружает PDF по ссылке → конвертирует в DOCX → извлекает расписание → сохраняет JSON по группам
"""

import os
import re
import json
import time
import logging
import argparse
import requests
from pathlib import Path
from datetime import date, timedelta
from docx import Document
from docx.table import Table

# ─── Настройки ──────────────────────────────────────────────────────────────

GROUPS_COUNT = 7
GROUP_PREFIX = "ВМ.О-ВЕТ.С-22-"
GROUPS = [f"{GROUP_PREFIX}{i}" for i in range(1, GROUPS_COUNT + 1)]

# Аббревиатуры дней → полное название и номер (ISO: пн=1)
DAY_MAP = {
    "ПНД": ("Понедельник", 1),
    "ВТР": ("Вторник",     2),
    "СРД": ("Среда",       3),
    "ЧТВ": ("Четверг",     4),
    "ПТН": ("Пятница",     5),
    "СБТ": ("Суббота",     6),
    "ВСК": ("Воскресенье", 7),
}
SKIP_DAYS = {"СБТ", "ВСК"}   # исключаем из расписания

MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая",    6: "июня",    7: "июля",  8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

PDF_FILENAME  = "schedule.pdf"
DOCX_FILENAME = "schedule_converted.docx"
OUTPUT_DIR    = "schedule_json"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)


# ─── Утилиты ────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Убирает лишние пробелы и артефакты bold-форматирования."""
    text = text.strip()
    text = re.sub(r"\s+", " ", text)
    return text


def get_cell_text(cell) -> str:
    """Собирает весь текст из параграфов ячейки."""
    parts = []
    for para in cell.paragraphs:
        parts.append(para.text)
    return clean_text(" ".join(parts))


def parse_day_date(raw: str):
    """
    Парсит строку вида 'ВТР 7/04' или 'ПНД 11/05'.
    Возвращает (abbr, day_n, month_n) или None.
    """
    raw = clean_text(raw)
    m = re.match(r"([А-ЯЁ]{3})\s+(\d{1,2})/(\d{2})", raw)
    if not m:
        return None
    abbr = m.group(1)
    return abbr, int(m.group(2)), int(m.group(3))


TIME_SLOT_ORDER = ("09.00", "10.45", "12.30", "14.30", "16.15")


def time_slot_index(slot: str) -> int:
    slot = (slot or "").strip()
    for i, prefix in enumerate(TIME_SLOT_ORDER):
        if slot.startswith(prefix):
            return i
    return -1


def should_advance_day(last_slot: str, new_slot: str) -> bool:
    """Новый слот раньше предыдущего → следующий учебный день (после разрыва страницы)."""
    li, ni = time_slot_index(last_slot), time_slot_index(new_slot)
    if li < 0 or ni < 0:
        return False
    return ni < li


def advance_weekday(abbr: str, day_n: int, month_n: int, year: int = 2026):
    """Следующий учебный день (пн–пт)."""
    try:
        d = date(year, month_n, day_n) + timedelta(days=1)
    except ValueError:
        return abbr, day_n, month_n
    while d.weekday() >= 5:
        d += timedelta(days=1)
    abbr_by_wd = ("ПНД", "ВТР", "СРД", "ЧТВ", "ПТН", "СБТ", "ВСК")
    return abbr_by_wd[d.weekday()], d.day, d.month


def date_range_str(days_in_week: list, year: int = 2026) -> str:
    """
    Принимает список (day_num, month_num) учебных дней недели.
    Возвращает строку диапазона, например '7–11 апреля 2026'.
    """
    if not days_in_week:
        return ""
    dates = sorted(days_in_week)
    first_day, first_month = dates[0]
    last_day,  last_month  = dates[-1]
    if first_month == last_month:
        return f"{first_day}–{last_day} {MONTHS_RU[first_month]} {year}"
    return (f"{first_day} {MONTHS_RU[first_month]} – "
            f"{last_day} {MONTHS_RU[last_month]} {year}")


# ─── Загрузка PDF ───────────────────────────────────────────────────────────

def download_pdf(url: str, dest: str = PDF_FILENAME) -> bool:
    """Скачивает PDF-файл по URL."""
    log.info(f"Загружаю PDF: {url}")
    try:
        r = requests.get(url, timeout=60, stream=True)
        r.raise_for_status()
        with open(dest, "wb") as f:
            for chunk in r.iter_content(chunk_size=65536):
                f.write(chunk)
        size_kb = Path(dest).stat().st_size // 1024
        log.info(f"Загружено {size_kb} КБ → {dest}")
        return True
    except Exception as e:
        log.error(f"Ошибка загрузки: {e}")
        return False


def convert_pdf_to_docx(pdf_path: str = PDF_FILENAME,
                         docx_path: str = DOCX_FILENAME) -> bool:
    """Конвертирует PDF в DOCX через pdf2docx."""
    try:
        from pdf2docx import Converter
        log.info(f"Конвертирую PDF → DOCX …")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        log.info(f"DOCX сохранён: {docx_path}")
        return True
    except Exception as e:
        log.error(f"Ошибка конвертации: {e}")
        return False


# ─── Парсинг DOCX ───────────────────────────────────────────────────────────

def find_header_row(table: Table):
    """
    Ищет строку-заголовок с 'Дни'/'Часы' и проверяет ширину таблицы.
    Возвращает индекс строки данных (после заголовков) или None.
    """
    for i, row in enumerate(table.rows):
        texts = [get_cell_text(c) for c in row.cells]
        # Ищем строку с 'Дни' и 'Часы'
        combined = " ".join(texts)
        if "Дни" in combined and "Часы" in combined:
            return i + 1   # данные начинаются со следующей строки
    return None


def parse_lesson_triple(cells, base_idx: int):
    """
    Извлекает тройку (дисциплина, преподаватель, аудитория)
    для одной группы, начиная с base_idx.
    """
    subj    = get_cell_text(cells[base_idx])     if base_idx   < len(cells) else ""
    teacher = get_cell_text(cells[base_idx + 1]) if base_idx+1 < len(cells) else ""
    room    = get_cell_text(cells[base_idx + 2]) if base_idx+2 < len(cells) else ""

    if not subj:
        return None
    return {
        "subject": subj,
        "teacher": teacher,
        "room":    room,
    }


def _process_row(row, state: dict, all_rows_data: list) -> None:
    """Одна строка расписания; state: current_day_info, last_time_slot."""
    cells = row.cells
    if not cells:
        return

    col0 = get_cell_text(cells[0])
    col1 = get_cell_text(cells[1]) if len(cells) > 1 else ""

    parsed = parse_day_date(col0)
    if parsed:
        state["current_day_info"] = parsed
        state["last_time_slot"] = None

    if state["current_day_info"] is None:
        return

    time_slot = col1.strip()
    if not re.match(r"(\d{2}\.\d{2}(?:-\d{2}\.\d{2})?)", time_slot):
        return

    if state["last_time_slot"] and should_advance_day(state["last_time_slot"], time_slot):
        abbr, day_n, month_n = state["current_day_info"]
        state["current_day_info"] = advance_weekday(abbr, day_n, month_n)

    abbr, day_n, month_n = state["current_day_info"]

    if abbr in SKIP_DAYS:
        state["last_time_slot"] = time_slot
        return

    full_row_text = " ".join(get_cell_text(c) for c in cells)
    if "самостоятельн" in full_row_text.lower():
        state["last_time_slot"] = time_slot
        return

    lessons = []
    for g in range(GROUPS_COUNT):
        lesson = parse_lesson_triple(cells, 2 + g * 3)
        lessons.append(lesson)

    all_rows_data.append((abbr, day_n, month_n, time_slot, lessons))
    state["last_time_slot"] = time_slot


def _parse_table(table: Table, all_rows_data: list, state: dict,
                 start_idx: int = 0, scan_header: bool = False) -> None:
    """
    Парсит таблицу и вложенные таблицы в ячейках (после разрыва страницы в Word).
    """
    if scan_header:
        hi = find_header_row(table)
        if hi is not None:
            start_idx = hi

    for row in table.rows[start_idx:]:
        for cell in row.cells:
            for nested in cell.tables:
                _parse_table(nested, all_rows_data, state, start_idx=0, scan_header=False)
        _process_row(row, state, all_rows_data)


def parse_tables(docx_path: str) -> list:
    """
    Главная функция парсинга.
    Возвращает список кортежей: (day_abbr, day_num, month_num, time, [lesson|None × 7])
    """
    doc = Document(docx_path)
    all_rows_data = []
    state = {"current_day_info": None, "last_time_slot": None}

    for t_idx, table in enumerate(doc.tables):
        _parse_table(
            table, all_rows_data, state,
            start_idx=0,
            scan_header=(t_idx == 0),
        )

    return all_rows_data


# ─── Группировка по неделям ─────────────────────────────────────────────────

def infer_year(month_n: int, prev_month: int | None, year: int) -> int:
    """Переход года если месяц уменьшился (крайний случай)."""
    if prev_month and month_n < prev_month and prev_month >= 11:
        return year + 1
    return year


def build_week_number(day_n: int, month_n: int, year: int = 2026) -> int:
    """ISO-номер недели для конкретной даты."""
    try:
        return date(year, month_n, day_n).isocalendar()[1]
    except ValueError:
        return 0


def group_by_weeks(all_rows_data: list, base_year: int = 2026) -> dict:
    """
    Группирует строки по группам и неделям.
    Возвращает:
    {
      "ВМ.О-ВЕТ.С-22-1": {
        week_iso: {
          "date_range": str,
          "days": {
            "7/04": {
              "day_name": "Вторник",
              "lessons": [ {"time":..., "subject":..., "teacher":..., "room":...}, ... ]
            }
          }
        }
      },
      ...
    }
    """
    # Инициализация
    result = {g: {} for g in GROUPS}
    year = base_year

    for abbr, day_n, month_n, time_slot, lessons in all_rows_data:
        week_iso = build_week_number(day_n, month_n, year)
        date_key = f"{day_n}/{month_n:02d}"  # "7/04"
        day_name = DAY_MAP.get(abbr, (abbr, 0))[0]

        for g_idx, lesson in enumerate(lessons):
            if lesson is None:
                continue
            group_name = GROUPS[g_idx]
            g_data = result[group_name]

            if week_iso not in g_data:
                g_data[week_iso] = {"date_range": "", "days": {}}

            week_data = g_data[week_iso]
            if date_key not in week_data["days"]:
                week_data["days"][date_key] = {
                    "day_name":   day_name,
                    "date":       date_key,
                    "day_n":      day_n,
                    "month_n":    month_n,
                    "lessons":    [],
                }

            week_data["days"][date_key]["lessons"].append({
                "time":    time_slot,
                "subject": lesson["subject"],
                "teacher": lesson["teacher"],
                "room":    lesson["room"],
            })

    # Формируем диапазоны дат для каждой недели
    for group_name, weeks in result.items():
        for week_iso, week_data in weeks.items():
            day_list = [(d["day_n"], d["month_n"])
                        for d in week_data["days"].values()]
            week_data["date_range"] = date_range_str(day_list, year=base_year)

    return result


# ─── Сериализация в JSON ─────────────────────────────────────────────────────

def export_to_json(grouped: dict, output_dir: str = OUTPUT_DIR):
    """Сохраняет отдельный JSON-файл для каждой группы."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    for group_name, weeks_dict in grouped.items():
        # Сортируем недели
        sorted_weeks = sorted(weeks_dict.items(), key=lambda x: x[0])

        output = {
            "group":       group_name,
            "generated":   time.strftime("%Y-%m-%dT%H:%M:%S"),
            "weeks": []
        }

        for week_iso, week_data in sorted_weeks:
            # Сортируем дни внутри недели по дате
            sorted_days = sorted(
                week_data["days"].values(),
                key=lambda d: (d["month_n"], d["day_n"])
            )
            for day in sorted_days:
                # Убираем служебные поля
                day.pop("day_n",   None)
                day.pop("month_n", None)
                # Сортируем пары по времени
                day["lessons"].sort(key=lambda l: l["time"])

            output["weeks"].append({
                "week":       week_iso,
                "date_range": week_data["date_range"],
                "days":       sorted_days,
            })

        # Имя файла: ВМ.О-ВЕТ.С-22-1 → VM_VET_S_22_1.json
        safe_name = re.sub(r"[^A-Za-zА-Яа-яЁё0-9]", "_", group_name)
        filepath = Path(output_dir) / f"{safe_name}.json"
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(output, f, ensure_ascii=False, indent=2)

        lessons_total = sum(
            len(d["lessons"])
            for w in output["weeks"]
            for d in w["days"]
        )
        log.info(f"  {group_name}: {len(output['weeks'])} нед., "
                 f"{sum(len(w['days']) for w in output['weeks'])} дней, "
                 f"{lessons_total} пар → {filepath.name}")

    log.info(f"Все файлы сохранены в '{output_dir}/'")


# ─── Основной цикл ───────────────────────────────────────────────────────────

def run_once(url: str, pdf_path: str, docx_path: str, output_dir: str):
    """Одна итерация: скачать → конвертировать → распарсить → сохранить."""
    if not download_pdf(url, pdf_path):
        return False
    if not convert_pdf_to_docx(pdf_path, docx_path):
        # Если конвертация не удалась — пробуем работать с существующим DOCX
        if not Path(docx_path).exists():
            log.error("DOCX-файл недоступен, пропускаем итерацию.")
            return False
        log.warning("Используем ранее конвертированный DOCX.")

    log.info("Парсю расписание …")
    raw_data = parse_tables(docx_path)
    log.info(f"Найдено строк с данными: {len(raw_data)}")

    grouped = group_by_weeks(raw_data)
    export_to_json(grouped, output_dir)
    return True


def main():
    parser = argparse.ArgumentParser(
        description="Парсер расписания: PDF → DOCX → JSON по группам"
    )
    parser.add_argument(
        "--url", "-u",
        help="Прямая ссылка на PDF-файл расписания (пропустите для ввода вручную)"
    )
    parser.add_argument(
        "--interval", "-i", type=int, default=0,
        help="Интервал автообновления в минутах (0 = однократный запуск)"
    )
    parser.add_argument(
        "--pdf",  default=PDF_FILENAME,  help=f"Путь для PDF  (по умолч. {PDF_FILENAME})"
    )
    parser.add_argument(
        "--docx", default=DOCX_FILENAME, help=f"Путь для DOCX (по умолч. {DOCX_FILENAME})"
    )
    parser.add_argument(
        "--output", "-o", default=OUTPUT_DIR,
        help=f"Каталог для JSON-файлов (по умолч. {OUTPUT_DIR})"
    )
    parser.add_argument(
        "--local-docx", "-l",
        help="Использовать готовый DOCX без загрузки PDF (для тестирования)"
    )
    args = parser.parse_args()

    # Если передан готовый DOCX — только парсим
    if args.local_docx:
        log.info(f"Режим: локальный DOCX-файл '{args.local_docx}'")
        raw_data = parse_tables(args.local_docx)
        log.info(f"Найдено строк с данными: {len(raw_data)}")
        grouped = group_by_weeks(raw_data)
        export_to_json(grouped, args.output)
        return

    # Получаем URL
    url = args.url
    if not url:
        url = input("Введите ссылку для скачивания PDF-файла расписания:\n> ").strip()
    if not url:
        log.error("URL не указан. Выход.")
        return

    if args.interval <= 0:
        # Однократный запуск
        run_once(url, args.pdf, args.docx, args.output)
    else:
        # Периодическое обновление
        log.info(f"Автообновление каждые {args.interval} мин. Ctrl+C для остановки.")
        iteration = 0
        while True:
            iteration += 1
            log.info(f"─── Итерация #{iteration} ───")
            run_once(url, args.pdf, args.docx, args.output)
            log.info(f"Следующее обновление через {args.interval} мин. …")
            try:
                time.sleep(args.interval * 60)
            except KeyboardInterrupt:
                log.info("Остановлено пользователем.")
                break


if __name__ == "__main__":
    main()
