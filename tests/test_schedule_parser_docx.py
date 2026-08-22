# -*- coding: utf-8 -*-
"""
Регрессионные тесты на конкретные баги парсера, которые уже один раз ловили
на реальном PDF колледжа:
  1. pdf2docx иногда прячет время пары во вложенную таблицу внутри ячейки
     вместо простого текста (get_time_text должен туда заглянуть).
  2. pdf2docx иногда схлопывает соседние ячейки в объединённую (merged) —
     парсер не должен падать и должен сохранять структуру строки.

Тесты строят синтетический .docx через python-docx (реальный PDF колледжа
не хранится в репозитории) и прогоняют его через настоящий parse_tables(),
а не через мок — так тест ловит регрессии и в самой логике работы с
python-docx API, а не только в бизнес-логике.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from docx import Document

from schedule_parser import (
    get_cell_text, get_time_text, find_header_row, parse_lesson_triple,
    parse_tables, GROUPS_COUNT,
)

N_COLS = 2 + GROUPS_COUNT * 3  # дата, время, затем (предмет, препод, ауд.) × 7 групп


@pytest.fixture
def tmp_docx_path(tmp_path):
    return str(tmp_path / "schedule_test.docx")


def _build_basic_table(doc):
    """Заголовок + одна строка с данными (только группа 0 заполнена)."""
    table = doc.add_table(rows=2, cols=N_COLS)
    table.cell(0, 0).text = "Дни"
    table.cell(0, 1).text = "Часы"
    table.cell(1, 0).text = "ВТР 7/04"
    table.cell(1, 1).text = "09.00-10.30"
    table.cell(1, 2).text = "Химия"
    table.cell(1, 3).text = "Иванов И.И."
    table.cell(1, 4).text = "УЛК-101"
    return table


def test_find_header_row_locates_data_start():
    doc = Document()
    table = _build_basic_table(doc)
    assert find_header_row(table) == 1  # данные начинаются со строки 1


def test_parse_lesson_triple_reads_subject_teacher_room():
    doc = Document()
    table = _build_basic_table(doc)
    row = table.rows[1]
    lesson = parse_lesson_triple(row.cells, base_idx=2)
    assert lesson == {"subject": "Химия", "teacher": "Иванов И.И.", "room": "УЛК-101"}


def test_parse_lesson_triple_empty_subject_returns_none():
    doc = Document()
    table = doc.add_table(rows=1, cols=5)
    lesson = parse_lesson_triple(table.rows[0].cells, base_idx=2)
    assert lesson is None


def test_full_parse_end_to_end_basic_table(tmp_docx_path):
    """Полный прогон через parse_tables() — сохраняем на диск и читаем заново,
    как это делает настоящий парсер (Document(docx_path))."""
    doc = Document()
    _build_basic_table(doc)
    doc.save(tmp_docx_path)

    rows = parse_tables(tmp_docx_path)
    assert len(rows) == 1
    abbr, day_n, month_n, time_slot, lessons = rows[0]
    assert (abbr, day_n, month_n, time_slot) == ("ВТР", 7, 4, "09.00-10.30")
    assert lessons[0] == {"subject": "Химия", "teacher": "Иванов И.И.", "room": "УЛК-101"}
    assert all(l is None for l in lessons[1:])  # остальные 6 групп пустые


# ─── Регрессия №1: время во вложенной таблице ────────────────────────────────

def test_get_time_text_falls_back_to_nested_table():
    """pdf2docx иногда кладёт время не текстом в ячейку, а вложенной таблицей
    1×1 внутри неё — get_time_text обязан достать значение оттуда."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    assert get_cell_text(cell) == ""  # прямого текста нет — это и есть баг-условие
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "12.30-14.00"
    assert get_time_text(cell) == "12.30-14.00"


def test_get_time_text_prefers_direct_text_over_nested():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "09.00-10.30"
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "16.15-17.45"  # не должно перекрыть прямой текст
    assert get_time_text(cell) == "09.00-10.30"


def test_full_parse_with_nested_time_table(tmp_docx_path):
    """Сквозной тест: та же регрессия, но через реальный parse_tables()."""
    doc = Document()
    table = _build_basic_table(doc)
    time_cell = table.cell(1, 1)
    time_cell.text = ""  # убираем прямой текст времени
    nested = time_cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "09.00-10.30"
    doc.save(tmp_docx_path)

    rows = parse_tables(tmp_docx_path)
    assert len(rows) == 1
    assert rows[0][3] == "09.00-10.30"  # время нашлось через вложенную таблицу


# ─── Регрессия №2: объединённые (merged) ячейки ──────────────────────────────

def test_merged_cells_do_not_crash_parser(tmp_docx_path):
    """pdf2docx иногда схлопывает пару соседних ячеек в одну merged-ячейку.
    python-docx в этом случае возвращает один и тот же Cell-объект для всех
    охваченных позиций в row.cells — важно, чтобы parse_tables() на этом
    не падал и не терял остальные строки таблицы."""
    doc = Document()
    table = _build_basic_table(doc)
    # объединяем "препод" и "аудиторию" первой группы (типичное место мёрджа)
    table.cell(1, 3).merge(table.cell(1, 4))
    doc.save(tmp_docx_path)

    rows = parse_tables(tmp_docx_path)  # не должно бросить исключение
    assert len(rows) == 1
    abbr, day_n, month_n, time_slot, lessons = rows[0]
    assert lessons[0]["subject"] == "Химия"  # предмет вне мёрджа — не пострадал
    assert len(lessons) == GROUPS_COUNT  # структура строки (7 групп) не нарушена


def test_merged_cells_across_full_row_still_parses_other_rows(tmp_docx_path):
    """Мёрдж в одной строке не должен ломать парсинг следующей за ней строки."""
    doc = Document()
    table = doc.add_table(rows=3, cols=N_COLS)
    table.cell(0, 0).text = "Дни"
    table.cell(0, 1).text = "Часы"
    table.cell(1, 0).text = "ВТР 7/04"
    table.cell(1, 1).text = "09.00-10.30"
    table.cell(1, 2).merge(table.cell(1, 3))  # мёрдж в первой строке данных
    table.cell(1, 2).text = "Химия"
    table.cell(2, 1).text = "10.45-12.15"
    table.cell(2, 2).text = "Физика"
    table.cell(2, 3).text = "Петров П.П."
    table.cell(2, 4).text = "УЛК-102"
    doc.save(tmp_docx_path)

    rows = parse_tables(tmp_docx_path)
    assert len(rows) == 2
    assert rows[1][3] == "10.45-12.15"
    assert rows[1][4][0]["subject"] == "Физика"
